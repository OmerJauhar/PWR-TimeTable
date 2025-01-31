import requests
import os
from dotenv import load_dotenv
from docx import Document  # Importing the Document class to read Word files

# Load environment variables from .env file
load_dotenv()

def read_word_file(file_path):
    """Read the content of a Word file and return it as a string."""
    doc = Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs if para.text])  # Join non-empty paragraphs

def generate_with_openrouter(prompt, model='anthropic/claude-2'):
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {os.environ['OPENROUTER_API_KEY']}",
            "Content-Type": "application/json"
        },
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}]
        }
    )
    
    return response.json()['choices'][0]['message']['content']

def save_output_to_file(output, file_name):
    """Save the generated output to a text file."""
    with open(file_name, 'w') as file:
        file.write(output)

def load_prompt_from_file(prompt_file_path):
    """Load the prompt from a text file and return it as a string."""
    with open(prompt_file_path, 'r') as file:
        return file.read()

def load_meow_file(meow_file_path):
    """Load the content from meow.docx and return it as a string."""
    doc = Document(meow_file_path)  # Use Document to read .docx file
    return '\n'.join([para.text for para in doc.paragraphs if para.text])  # Join non-empty paragraphs

# Example usage
word_file_path = "/home/oj/repos/PWR-TimeTable/script /meow.docx"  # Update this to your actual file name
file_content = read_word_file(word_file_path)  # Read the content of the Word file
prompt_file_path = "prompt.txt"  # Specify the path to the prompt file
prompt = load_prompt_from_file(prompt_file_path)  # Load the prompt from the file

# New code to load meow.txt and concatenate with the prompt
meow_file_path = "meow.docx"  # Specify the path to the meow file
meow_content = load_meow_file(meow_file_path)  # Load the content from meow.txt
prompt += "\n" + meow_content  # Concatenate meow content with the prompt

result = generate_with_openrouter(prompt)  # Generate the output

# Save the result to a .txt file
output_file_path = "output.txt"  # Specify the output file name
save_output_to_file(result, output_file_path)  # Save the output to a text file
print(f"Output saved to {output_file_path}")